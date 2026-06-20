import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

md_path = r"C:\Users\WAZIR AMMAR HAIDER\Desktop\fyp\HCI Assignment 2\APCRE_HCI_Assignment_Report.md"
image_dir = r"C:\Users\WAZIR AMMAR HAIDER\Desktop\fyp\HCI Assignment 2"
output_path = r"C:\Users\WAZIR AMMAR HAIDER\Desktop\fyp\HCI Assignment 2\APCRE_HCI_Assignment_Report.docx"

def set_run_font(run, name="Calibri", size_pt=11, color_rgb=(0, 0, 0), bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, name="Calibri Light", size_pt=18, color_rgb=(47, 85, 151), bold=True)
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, name="Calibri Light", size_pt=14, color_rgb=(47, 85, 151), bold=True)
    elif level == 3:
        run = p.add_run(text)
        set_run_font(run, name="Calibri Light", size_pt=12, color_rgb=(0, 0, 0), bold=True, italic=True)
    return p

def add_styled_paragraph(doc, text, style_type="body"):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    
    if style_type == "bullet":
        p.paragraph_format.left_indent = Inches(0.25)
        run_prefix = p.add_run("•  ")
        set_run_font(run_prefix, name="Calibri", size_pt=11)
    elif style_type == "number":
        p.paragraph_format.left_indent = Inches(0.25)
        # Prefix added inside loop
        
    # Parse inline formatting (**bold**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            run = p.add_run(inner)
            set_run_font(run, name="Calibri", size_pt=11, bold=True)
        else:
            if part:
                run = p.add_run(part)
                set_run_font(run, name="Calibri", size_pt=11)
    return p

def create_table(doc, rows_data):
    # rows_data is list of lists
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    for r_idx, row in enumerate(rows_data):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = "" # Clear default text
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            
            # Formatting: Bold headers
            if r_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                set_run_font(run, name="Calibri", size_pt=10, bold=True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Parse inline bold in cell values
                parts = re.split(r'(\*\*.*?\*\*)', val)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        inner = part[2:-2]
                        run = p.add_run(inner)
                        set_run_font(run, name="Calibri", size_pt=9.5, bold=True)
                    else:
                        if part:
                            run = p.add_run(part)
                            set_run_font(run, name="Calibri", size_pt=9.5)
                            
    # Add a spacing paragraph after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)

def build_docx():
    print(f"Reading Markdown file...")
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    doc = Document()
    
    # Page Setup Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    lines = content.splitlines()
    in_table = False
    table_lines = []
    
    # Title Header Section (centered title block)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("UNIVERSITY OF ENGINEERING AND TECHNOLOGY, TAXILA\n")
    set_run_font(run_title, name="Calibri Light", size_pt=14, color_rgb=(47, 85, 151), bold=True)
    run_sub = title_p.add_run("FACULTY OF TELECOMMUNICATION AND INFORMATION ENGINEERING\nSOFTWARE ENGINEERING DEPARTMENT\n\n")
    set_run_font(run_sub, name="Calibri", size_pt=11, color_rgb=(0, 0, 0), bold=True)
    
    # Process remaining lines
    for line in lines:
        stripped = line.strip()
        
        # Skip top header definitions that we printed centered
        if any(h in stripped for h in ["UNIVERSITY OF ENGINEERING", "TELECOMMUNICATION AND", "SOFTWARE ENGINEERING DEPARTMENT"]):
            continue
            
        # Parse Tables
        if stripped.startswith("|"):
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            # Table finished, build it
            valid_rows = []
            for r in table_lines:
                cells = [c.strip() for c in r.split("|")[1:-1]]
                if cells and all(re.match(r'^:?-+:?$', c) for c in cells):
                    continue
                valid_rows.append(cells)
            if valid_rows:
                create_table(doc, valid_rows)
            table_lines = []
            in_table = False
            
        if not stripped:
            continue
            
        # Headings
        if stripped.startswith("# "):
            add_styled_heading(doc, stripped[2:], 1)
        elif stripped.startswith("## "):
            add_styled_heading(doc, stripped[3:], 2)
        elif stripped.startswith("### "):
            add_styled_heading(doc, stripped[4:], 3)
            
        # Lists
        elif stripped.startswith("* ") or stripped.startswith("- "):
            add_styled_paragraph(doc, stripped[2:], "bullet")
        elif re.match(r'^\d+\.\s+', stripped):
            num_match = re.match(r'^(\d+\.\s+)', stripped)
            prefix = num_match.group(1)
            text = stripped[len(prefix):]
            p = add_styled_paragraph(doc, text, "number")
            # insert numbers prefix at start of paragraph
            p.paragraph_format.left_indent = Inches(0.25)
            r_num = p.add_run()
            r_num.text = prefix + " "
            set_run_font(r_num, name="Calibri", size_pt=11)
            # Reorder run to front
            p.runs.insert(0, p.runs.pop())
            
        # Mermaid code triggers skip
        elif stripped.startswith("```"):
            continue
            
        # Body text
        else:
            add_styled_paragraph(doc, stripped, "body")
            
            # IMAGE EMBED TRIGGER 1: after Low-Fidelity sketch reference
            if "apcre_low_fi_sketch.png" in stripped:
                img_path = os.path.join(image_dir, "apcre_low_fi_sketch.png")
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(12)
                    p_img.paragraph_format.space_after = Pt(4)
                    run_img = p_img.add_run()
                    run_img.add_picture(img_path, width=Inches(5.5))
                    
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_after = Pt(12)
                    run_cap = p_cap.add_run("Figure 2.1: Low-Fidelity Paper Prototype Sketch (APCRE IDE Dashboard Layout)")
                    set_run_font(run_cap, name="Calibri", size_pt=9.5, color_rgb=(89, 89, 89), italic=True)
                    
            # IMAGE EMBED TRIGGER 2: after High-Fidelity mockup reference
            elif "apcre_high_fi_mockup.png" in stripped:
                img_path = os.path.join(image_dir, "apcre_high_fi_mockup.png")
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(12)
                    p_img.paragraph_format.space_after = Pt(4)
                    run_img = p_img.add_run()
                    run_img.add_picture(img_path, width=Inches(5.5))
                    
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_after = Pt(12)
                    run_cap = p_cap.add_run("Figure 3.1: High-Fidelity UI Mockup (APCRE Dark-Mode IDE Dashboard)")
                    set_run_font(run_cap, name="Calibri", size_pt=9.5, color_rgb=(89, 89, 89), italic=True)
                    
    # Clean up and save
    try:
        doc.save(output_path)
        print(f"Saved highly styled docx document with embedded pictures to: {output_path}")
    except PermissionError:
        backup_path = r"C:\Users\WAZIR AMMAR HAIDER\Desktop\fyp\HCI Assignment 2\APCRE_HCI_Assignment_Report_Final.docx"
        doc.save(backup_path)
        print(f"[NOTE] Original file was locked. Saved the final styled docx with embedded pictures to: {backup_path}")

if __name__ == "__main__":
    build_docx()
